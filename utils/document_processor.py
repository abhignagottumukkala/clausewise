"""
Document processing utilities for ClauseWise
Handles PDF, DOCX, TXT, and other file processing
"""

import PyPDF2
import docx
import os
import io
from typing import Optional, Dict, Any
import streamlit as st
import base64


class DocumentProcessor:
    """Handles document processing for various file formats"""
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.txt', '.rtf', '.doc', '.odt']
    
    def extract_text_from_pdf(self, pdf_file) -> str:
        """Extract text from PDF file"""
        try:
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            return text.strip()
        except Exception as e:
            st.error(f"Error processing PDF: {str(e)}")
            return ""
    
    def extract_text_from_docx(self, docx_file) -> str:
        """Extract text from DOCX file"""
        try:
            doc = docx.Document(docx_file)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text.strip()
        except Exception as e:
            st.error(f"Error processing DOCX: {str(e)}")
            return ""
    
    def extract_text_from_txt(self, txt_file) -> str:
        """Extract text from TXT file"""
        try:
            text = txt_file.read().decode('utf-8')
            return text.strip()
        except Exception as e:
            st.error(f"Error processing TXT: {str(e)}")
            return ""
    
    def extract_text_from_rtf(self, rtf_file) -> str:
        """Extract text from RTF file (basic implementation)"""
        try:
            # Basic RTF text extraction
            content = rtf_file.read().decode('utf-8', errors='ignore')
            # Remove RTF formatting tags (simplified)
            import re
            # Remove RTF control words
            content = re.sub(r'\\[a-z]+\d*', '', content)
            # Remove braces
            content = content.replace('{', '').replace('}', '')
            # Remove extra whitespace
            content = re.sub(r'\s+', ' ', content)
            return content.strip()
        except Exception as e:
            st.error(f"Error processing RTF: {str(e)}")
            return ""
    
    def process_document(self, uploaded_file) -> Dict[str, Any]:
        """Process uploaded document and extract text"""
        if uploaded_file is None:
            return {"success": False, "text": "", "error": "No file uploaded"}
        
        file_extension = os.path.splitext(uploaded_file.name)[1].lower()
        
        if file_extension not in self.supported_formats:
            return {
                "success": False, 
                "text": "", 
                "error": f"Unsupported file format. Supported formats: {', '.join(self.supported_formats)}"
            }
        
        try:
            if file_extension == '.pdf':
                text = self.extract_text_from_pdf(uploaded_file)
            elif file_extension == '.docx':
                text = self.extract_text_from_docx(uploaded_file)
            elif file_extension == '.txt':
                text = self.extract_text_from_txt(uploaded_file)
            elif file_extension == '.rtf':
                text = self.extract_text_from_rtf(uploaded_file)
            elif file_extension in ['.doc', '.odt']:
                # For .doc and .odt files, try to extract as text
                text = self.extract_text_from_txt(uploaded_file)
            else:
                return {"success": False, "text": "", "error": "Unsupported file format"}
            
            if not text:
                return {"success": False, "text": "", "error": "No text extracted from document"}
            
            return {
                "success": True,
                "text": text,
                "filename": uploaded_file.name,
                "file_size": len(text),
                "file_type": file_extension
            }
            
        except Exception as e:
            return {"success": False, "text": "", "error": f"Error processing document: {str(e)}"}
    
    def process_uploaded_file(self, uploaded_file) -> Dict[str, Any]:
        """Process uploaded file - alias for process_document for compatibility"""
        return self.process_document(uploaded_file)
    
    def split_into_clauses(self, text: str, max_length: int = 1000) -> list:
        """Split document text into manageable clauses"""
        # Enhanced clause splitting based on common legal document patterns
        clause_markers = [
            "WHEREAS", "NOW, THEREFORE", "IN WITNESS WHEREOF", "SECTION", 
            "ARTICLE", "CLAUSE", "PROVIDED", "PROVIDED THAT", "FURTHER",
            "ADDITIONALLY", "MOREOVER", "FURTHERMORE", "IN ADDITION",
            "THEREFORE", "HEREBY", "AGREES", "AGREED", "PARTIES",
            "DEFINITIONS", "SCOPE", "TERM", "TERMINATION", "LIABILITY",
            "INDEMNIFICATION", "CONFIDENTIALITY", "NON-DISCLOSURE",
            "PAYMENT", "COMPENSATION", "BENEFITS", "DUTIES", "OBLIGATIONS",
            "REPRESENTATIONS", "WARRANTIES", "COVENANTS", "CONDITIONS",
            "DEFAULT", "BREACH", "REMEDIES", "DISPUTE", "ARBITRATION",
            "GOVERNING LAW", "JURISDICTION", "AMENDMENT", "WAIVER",
            "SEVERABILITY", "ENTIRE AGREEMENT", "FORCE MAJEURE",
            "NOTICES", "ASSIGNMENT", "SUCCESSORS", "COUNTERPARTS"
        ]
        
        # Additional patterns for numbered sections
        numbered_patterns = [
            r'^\d+\.',  # 1. 2. 3.
            r'^[a-z]\)',  # a) b) c)
            r'^[A-Z]\.',  # A. B. C.
            r'^\([a-z]\)',  # (a) (b) (c)
            r'^\([A-Z]\)',  # (A) (B) (C)
        ]
        
        import re
        
        clauses = []
        current_clause = ""
        
        lines = text.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            # Check if line starts with a clause marker
            is_clause_start = any(line.upper().startswith(marker) for marker in clause_markers)
            
            # Check for numbered patterns
            is_numbered_start = any(re.match(pattern, line) for pattern in numbered_patterns)
            
            # Check for paragraph breaks (double line breaks)
            is_paragraph_break = len(line) < 50 and line.isupper()
            
            if (is_clause_start or is_numbered_start or is_paragraph_break) and current_clause:
                clauses.append(current_clause.strip())
                current_clause = line
            else:
                current_clause += " " + line if current_clause else line
            
            # If clause is getting too long, split it
            if len(current_clause) > max_length:
                # Try to split at sentence boundaries
                sentences = re.split(r'[.!?]+', current_clause)
                if len(sentences) > 1:
                    # Split at the middle sentence
                    mid_point = len(sentences) // 2
                    first_part = '. '.join(sentences[:mid_point]) + '.'
                    second_part = '. '.join(sentences[mid_point:])
                    
                    clauses.append(first_part.strip())
                    current_clause = second_part
                else:
                    clauses.append(current_clause.strip())
                    current_clause = ""
        
        if current_clause:
            clauses.append(current_clause.strip())
        
        # Filter out very short clauses and clean up
        filtered_clauses = []
        for clause in clauses:
            clause = clause.strip()
            if len(clause) > 20:  # Minimum length for a meaningful clause
                # Clean up extra whitespace
                clause = re.sub(r'\s+', ' ', clause)
                filtered_clauses.append(clause)
        
        return filtered_clauses
    
    def get_document_stats(self, text: str) -> Dict[str, Any]:
        """Get basic statistics about the document"""
        words = text.split()
        sentences = text.split('.')
        paragraphs = text.split('\n\n')
        
        return {
            "word_count": len(words),
            "sentence_count": len([s for s in sentences if s.strip()]),
            "paragraph_count": len([p for p in paragraphs if p.strip()]),
            "character_count": len(text),
            "estimated_reading_time": len(words) // 200  # Average reading speed
        }
    
    def create_download_link(self, text: str, filename: str, file_type: str = "txt") -> str:
        """Create a download link for the processed text"""
        try:
            if file_type == "txt":
                # Create text file
                b64 = base64.b64encode(text.encode()).decode()
                return f'<a href="data:file/txt;base64,{b64}" download="{filename}" target="_blank">Download {filename}</a>'
            
            elif file_type == "docx":
                # Create DOCX file
                from docx import Document
                doc = Document()
                doc.add_paragraph(text)
                
                # Save to bytes
                docx_bytes = io.BytesIO()
                doc.save(docx_bytes)
                docx_bytes.seek(0)
                
                b64 = base64.b64encode(docx_bytes.read()).decode()
                return f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="{filename}.docx" target="_blank">Download {filename}.docx</a>'
            
            elif file_type == "pdf":
                # For PDF, we'll create a simple text-based PDF
                try:
                    from reportlab.pdfgen import canvas
                    from reportlab.lib.pagesizes import letter
                    from reportlab.lib.units import inch
                    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
                    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
                    from reportlab.lib.enums import TA_LEFT, TA_CENTER
                    
                    # Create PDF in memory
                    buffer = io.BytesIO()
                    doc = SimpleDocTemplate(buffer, pagesize=letter)
                    styles = getSampleStyleSheet()
                    
                    # Create custom style for the title
                    title_style = ParagraphStyle(
                        'CustomTitle',
                        parent=styles['Heading1'],
                        fontSize=16,
                        spaceAfter=30,
                        alignment=TA_CENTER
                    )
                    
                    # Create custom style for body text
                    body_style = ParagraphStyle(
                        'CustomBody',
                        parent=styles['Normal'],
                        fontSize=11,
                        spaceAfter=12,
                        alignment=TA_LEFT
                    )
                    
                    # Build the PDF content
                    story = []
                    
                    # Add title
                    story.append(Paragraph("ClauseWise - Document Summary", title_style))
                    story.append(Spacer(1, 20))
                    
                    # Add the text content
                    lines = text.split('\n')
                    for line in lines:
                        if line.strip():
                            story.append(Paragraph(line, body_style))
                    
                    # Build the PDF
                    doc.build(story)
                    buffer.seek(0)
                    
                    # Create download link
                    b64 = base64.b64encode(buffer.read()).decode()
                    return f'<a href="data:application/pdf;base64,{b64}" download="{filename}.pdf" target="_blank">Download {filename}.pdf</a>'
                    
                except ImportError:
                    # Fallback to text if reportlab is not available
                    st.warning("PDF generation requires reportlab. Installing simplified version...")
                    b64 = base64.b64encode(text.encode()).decode()
                    return f'<a href="data:file/txt;base64,{b64}" download="{filename}.txt" target="_blank">Download {filename}.txt</a>'
            
            else:
                # Default to text
                b64 = base64.b64encode(text.encode()).decode()
                return f'<a href="data:file/txt;base64,{b64}" download="{filename}.txt" target="_blank">Download {filename}.txt</a>'
                
        except Exception as e:
            st.error(f"Error creating download link: {str(e)}")
            return ""
    
    def get_supported_formats_display(self) -> str:
        """Get a user-friendly list of supported formats"""
        format_names = {
            '.pdf': 'PDF Documents',
            '.docx': 'Word Documents (DOCX)',
            '.txt': 'Text Files',
            '.rtf': 'Rich Text Format',
            '.doc': 'Word Documents (DOC)',
            '.odt': 'OpenDocument Text'
        }
        
        return ", ".join([format_names.get(fmt, fmt) for fmt in self.supported_formats])
    
    def summarize_document(self, text: str) -> str:
        """Create a summary of the legal document"""
        # Split into sentences
        sentences = text.split('.')
        
        # Extract key information
        summary_parts = []
        
        # Document type detection
        text_lower = text.lower()
        if any(term in text_lower for term in ['confidential', 'non-disclosure', 'nda']):
            doc_type = "Non-Disclosure Agreement (NDA)"
        elif any(term in text_lower for term in ['employment', 'employee', 'hire']):
            doc_type = "Employment Contract"
        elif any(term in text_lower for term in ['lease', 'rent', 'tenant']):
            doc_type = "Lease Agreement"
        elif any(term in text_lower for term in ['service', 'vendor', 'provider']):
            doc_type = "Service Agreement"
        else:
            doc_type = "Legal Contract"
        
        summary_parts.append(f"Document Type: {doc_type}")
        summary_parts.append("")
        
        # Key parties
        parties = []
        for sentence in sentences[:10]:  # Check first 10 sentences
            if any(word in sentence.lower() for word in ['corporation', 'company', 'inc', 'llc', 'ltd']):
                parties.append(sentence.strip())
        
        if parties:
            summary_parts.append("Key Parties:")
            summary_parts.extend([f"- {party}" for party in parties[:3]])
            summary_parts.append("")
        
        # Key terms and conditions
        key_terms = []
        for sentence in sentences:
            if any(term in sentence.lower() for term in ['shall', 'must', 'agree', 'obligated', 'liability', 'termination']):
                key_terms.append(sentence.strip())
        
        if key_terms:
            summary_parts.append("Key Terms and Conditions:")
            summary_parts.extend([f"- {term}" for term in key_terms[:5]])
            summary_parts.append("")
        
        # Document statistics
        stats = self.get_document_stats(text)
        summary_parts.append("Document Statistics:")
        summary_parts.append(f"- Word Count: {stats['word_count']:,}")
        summary_parts.append(f"- Estimated Reading Time: {stats['estimated_reading_time']} minutes")
        summary_parts.append("")
        
        # Main points
        summary_parts.append("Main Points:")
        summary_parts.append("- This document contains legal terms and conditions")
        summary_parts.append("- Review all clauses carefully before signing")
        summary_parts.append("- Consider consulting with legal counsel")
        
        return "\n".join(summary_parts) 